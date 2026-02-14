"""
RAG-хранилище на ChromaDB.
Загрузка документов, чанкинг, эмбеддинги OpenAI, индексация.
"""

import hashlib
import logging
import re
from pathlib import Path
from typing import Iterator

from openai import OpenAI

logger = logging.getLogger(__name__)

DATA_DIR = Path("data")
CHUNK_SIZE = 500
CHUNK_OVERLAP = 50
EMBEDDING_MODEL = "text-embedding-3-small"


def _chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Разбивает текст на чанки с перекрытием."""
    if not text or not text.strip():
        return []
    text = text.strip()
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        if end < len(text):
            last_space = chunk.rfind(" ")
            if last_space > chunk_size // 2:
                chunk = chunk[: last_space + 1]
                end = start + last_space + 1
        chunk = chunk.strip()
        if chunk:
            chunks.append(chunk)
        start = end - overlap if end < len(text) else len(text)
    return chunks


def _load_txt(path: Path) -> str:
    """Загружает текстовый файл."""
    return path.read_text(encoding="utf-8", errors="replace")


def _load_pdf(path: Path) -> str:
    """Загружает PDF и извлекает текст."""
    from pypdf import PdfReader
    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n\n".join(parts)


def _load_xlsx(path: Path) -> str:
    """Загружает XLSX и извлекает текст из ячеек."""
    from openpyxl import load_workbook
    wb = load_workbook(path, read_only=True, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join(str(c) if c is not None else "" for c in row)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def _load_docx(path: Path) -> str:
    """Загружает DOCX и извлекает текст."""
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text)
            if row_text.strip():
                parts.append(row_text)
    return "\n\n".join(parts)


def load_document(path: Path) -> str | None:
    """Загружает документ по расширению. Возвращает текст или None."""
    suf = path.suffix.lower()
    try:
        if suf in (".txt", ".text", ".md"):
            return _load_txt(path)
        if suf == ".pdf":
            return _load_pdf(path)
        if suf in (".xlsx", ".xls"):
            return _load_xlsx(path)
        if suf == ".docx":
            return _load_docx(path)
    except Exception as e:
        logger.exception("Ошибка загрузки %s: %s", path, e)
        return None
    return None


def get_embedding(client: OpenAI, text: str) -> list[float]:
    """Получает эмбеддинг через OpenAI."""
    resp = client.embeddings.create(model=EMBEDDING_MODEL, input=text.strip())
    return resp.data[0].embedding


class RAGStore:
    """Хранилище RAG на ChromaDB."""

    def __init__(self, persist_directory: str | Path = "chroma_db", api_key: str | None = None):
        self.persist_dir = Path(persist_directory)
        self.persist_dir.mkdir(parents=True, exist_ok=True)
        self.client = OpenAI(api_key=api_key)
        self._collection = None

    def _get_collection(self):
        import chromadb
        from chromadb.config import Settings
        client = chromadb.PersistentClient(path=str(self.persist_dir), settings=Settings(anonymized_telemetry=False))
        return client.get_or_create_collection(
            name="imagebot_rag",
            metadata={"description": "RAG collection for ImageBot"},
        )

    @property
    def collection(self):
        if self._collection is None:
            self._collection = self._get_collection()
        return self._collection

    def index_documents(self, data_dir: Path | None = None) -> dict[str, int]:
        """
        Индексирует документы из data_dir.
        Возвращает словарь {источник: количество чанков}.
        """
        data_dir = data_dir or DATA_DIR
        data_dir = Path(data_dir)
        if not data_dir.exists():
            data_dir.mkdir(parents=True, exist_ok=True)
            return {}

        exts = {".txt", ".text", ".md", ".pdf", ".xlsx", ".xls", ".docx"}
        files = [f for f in data_dir.rglob("*") if f.is_file() and f.suffix.lower() in exts]

        if not files:
            logger.info("Нет документов для индексации в %s", data_dir)
            return {}

        # Собираем все чанки
        all_ids: list[str] = []
        all_texts: list[str] = []
        all_metadatas: list[dict] = []
        source_counts: dict[str, int] = {}

        for path in sorted(files):
            text = load_document(path)
            if not text or not text.strip():
                continue
            source = str(path.relative_to(data_dir))
            chunks = _chunk_text(text)
            if not chunks:
                continue
            base_id = hashlib.md5(source.encode()).hexdigest()[:12]
            for i, chunk in enumerate(chunks):
                doc_id = f"{base_id}_{i}"
                all_ids.append(doc_id)
                all_texts.append(chunk)
                all_metadatas.append({"source": source})
            source_counts[source] = len(chunks)

        if not all_texts:
            return {}

        # Эмбеддинги батчами (OpenAI лимит ~8000 токенов на запрос, ~3000 слов)
        batch_size = 50
        all_embeddings: list[list[float]] = []

        for i in range(0, len(all_texts), batch_size):
            batch = all_texts[i : i + batch_size]
            emb = self.client.embeddings.create(model=EMBEDDING_MODEL, input=batch)
            all_embeddings.extend([e.embedding for e in emb.data])

        self.collection.upsert(
            ids=all_ids,
            embeddings=all_embeddings,
            documents=all_texts,
            metadatas=all_metadatas,
        )

        logger.info("Проиндексировано %d чанков из %d файлов", len(all_ids), len(source_counts))
        return source_counts

    def list_sources(self) -> list[str]:
        """Возвращает список уникальных источников в коллекции."""
        try:
            res = self.collection.get(include=["metadatas"])
            if not res or not res.get("metadatas"):
                return []
            sources = set()
            for m in res["metadatas"]:
                if m and "source" in m:
                    sources.add(m["source"])
            return sorted(sources)
        except Exception as e:
            logger.warning("Ошибка list_sources: %s", e)
            return []

    def delete_source(self, source: str, data_dir: Path | None = None) -> int:
        """
        Удаляет источник из ChromaDB и файл из data_dir (если существует).
        source: имя источника (как в list_sources, например "doc.pdf" или "folder/doc.pdf").
        Возвращает количество удалённых чанков.
        """
        source = source.strip()
        if not source:
            raise ValueError("Имя источника не может быть пустым")

        sources = self.list_sources()
        if source not in sources:
            raise ValueError(
                f"Источник '{source}' не найден. Используйте /rag_list для списка."
            )

        res = self.collection.get(
            where={"source": source},
            include=["metadatas"],
        )
        ids_to_delete = res.get("ids", [])
        if ids_to_delete:
            self.collection.delete(ids=ids_to_delete)
        count = len(ids_to_delete)

        data_dir = Path(data_dir or DATA_DIR)
        file_path = (data_dir / source).resolve()
        data_resolved = data_dir.resolve()
        if file_path.exists() and file_path.is_file():
            try:
                file_path.relative_to(data_resolved)
            except ValueError:
                logger.warning("Игнорируем удаление файла вне data/: %s", source)
            else:
                file_path.unlink()
                logger.info("Удалён файл %s", file_path)

        logger.info("Удалён источник %s: %d чанков", source, count)
        return count

    def query(self, query_text: str, n_results: int = 5) -> list[tuple[str, str, float]]:
        """
        Поиск по хранилищу. Возвращает [(document, source, distance), ...]
        """
        query_embedding = get_embedding(self.client, query_text)
        res = self.collection.query(
            query_embeddings=[query_embedding],
            n_results=n_results,
            include=["documents", "metadatas", "distances"],
        )
        if not res or not res.get("documents") or not res["documents"][0]:
            return []
        docs = res["documents"][0]
        metas = res.get("metadatas", [[]])[0] or [{}] * len(docs)
        dists = res.get("distances", [[]])[0] or [0.0] * len(docs)
        return [
            (doc, (m.get("source", "") if m else ""), d)
            for doc, m, d in zip(docs, metas, dists)
        ]
